# ------------------- Core Imports -------------------
import os, json, random, string, re, asyncio, io
import urllib.parse
from collections import Counter

# ------------------- External Libraries -------------------
import torch
import io
from io import BytesIO
import matplotlib.pyplot as plt
import fitz
import requests
import numpy as np
import pandas as pd
import streamlit as st
import base64
import streamlit as st

# File uploader widget for image upload
from db_manager import insert_candidate, get_top_domains_by_score

    
from PIL import Image
from pdf2image import convert_from_path
from dotenv import load_dotenv
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ------------------- Langchain & Embeddings -------------------
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from llm_manager import call_llm

from pydantic import BaseModel

# Set page config


# CSS Customization
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Orbitron', sans-serif;
        background-color: #0b0c10;
        color: #c5c6c7;
        scroll-behavior: smooth;
    }

    /* ---------- SCROLLBAR ---------- */
    ::-webkit-scrollbar {
        width: 8px;
    }
    ::-webkit-scrollbar-track {
        background: #1f2833;
    }
    ::-webkit-scrollbar-thumb {
        background: #00ffff;
        border-radius: 4px;
    }

    /* ---------- BANNER ---------- */
    .banner-container {
        width: 100%;
        height: 80px;
        background: linear-gradient(90deg, #000428, #004e92);
        border-bottom: 2px solid cyan;
        overflow: hidden;
        display: flex;
        align-items: center;
        justify-content: flex-start;
        position: relative;
        margin-bottom: 20px;
    }

    .pulse-bar {
        position: absolute;
        display: flex;
        align-items: center;
        font-size: 22px;
        font-weight: bold;
        color: #00ffff;
        white-space: nowrap;
        animation: glideIn 12s linear infinite;
        text-shadow: 0 0 10px #00ffff;
    }

    .pulse-bar .bar {
        width: 10px;
        height: 30px;
        margin-right: 10px;
        background: #00ffff;
        box-shadow: 0 0 8px cyan;
        animation: pulse 1s ease-in-out infinite;
    }

    @keyframes glideIn {
        0% { left: -50%; opacity: 0; }
        10% { opacity: 1; }
        90% { opacity: 1; }
        100% { left: 110%; opacity: 0; }
    }

    @keyframes pulse {
        0%, 100% {
            height: 20px;
            background-color: #00ffff;
        }
        50% {
            height: 40px;
            background-color: #ff00ff;
        }
    }

    /* ---------- HEADER ---------- */
    .header {
        font-size: 28px;
        font-weight: bold;
        text-align: center;
        text-transform: uppercase;
        letter-spacing: 2px;
        padding: 12px 0;
        animation: glowPulse 3s ease-in-out infinite;
        text-shadow: 0px 0px 10px #00ffff;
    }

    @keyframes glowPulse {
        0%, 100% {
            color: #00ffff;
            text-shadow: 0 0 10px #00ffff, 0 0 20px #00ffff;
        }
        50% {
            color: #ff00ff;
            text-shadow: 0 0 20px #ff00ff, 0 0 30px #ff00ff;
        }
    }

    /* ---------- FILE UPLOADER ---------- */
    .stFileUploader > div > div {
        border: 2px solid #00ffff;
        border-radius: 10px;
        background-color: rgba(0, 255, 255, 0.05);
        padding: 12px;
        box-shadow: 0 0 15px rgba(0,255,255,0.4);
        transition: box-shadow 0.3s ease-in-out;
    }
    .stFileUploader > div > div:hover {
        box-shadow: 0 0 25px rgba(0,255,255,0.8);
    }

    /* ---------- BUTTONS ---------- */
    .stButton > button {
        background: linear-gradient(45deg, #ff0080, #00bfff);
        color: white;
        font-size: 16px;
        font-weight: bold;
        border: none;
        border-radius: 8px;
        padding: 10px 20px;
        text-transform: uppercase;
        box-shadow: 0px 0px 12px #00ffff;
        transition: all 0.3s ease-in-out;
    }
    .stButton > button:hover {
        transform: scale(1.05);
        box-shadow: 0px 0px 24px #ff00ff;
        background: linear-gradient(45deg, #ff00aa, #00ffff);
    }

    /* ---------- CHAT MESSAGES ---------- */
    .stChatMessage {
        font-size: 18px;
        background: #1e293b;
        padding: 14px;
        border-radius: 10px;
        border: 2px solid #00ffff;
        color: #ccffff;
        text-shadow: 0px 0px 6px #00ffff;
        animation: glow 1.5s ease-in-out infinite alternate;
    }

    /* ---------- INPUTS ---------- */
    .stTextInput > div > input,
    .stTextArea > div > textarea {
        background-color: #1f2833;
        color: #00ffff;
        border: 1px solid #00ffff;
        border-radius: 6px;
        padding: 10px;
        box-shadow: 0 0 10px rgba(0,255,255,0.3);
    }

    /* ---------- METRICS ---------- */
    .stMetric {
        background-color: #0f172a;
        border: 1px solid #00ffff;
        border-radius: 10px;
        padding: 15px;
        box-shadow: 0 0 10px rgba(0,255,255,0.5);
        text-align: center;
    }

    /* ---------- MOBILE ---------- */
    @media (max-width: 768px) {
        .pulse-bar {
            font-size: 16px;
        }
        .header {
            font-size: 20px;
        }
    }
    </style>

    <!-- Banner -->
    <div class="banner-container">
        <div class="pulse-bar">
            <div class="bar"></div>
            <div>LEXIBOT - Elevate Your Resume Analysis</div>
        </div>
    </div>

    <!-- Header -->
    <div class="header">💼 LEXIBOT - AI ETHICAL RESUME ANALYZER</div>
    """,
    unsafe_allow_html=True
)

# Load environment variables
# ------------------- Core Setup -------------------


# Load environment variables
load_dotenv()

# Detect Device
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
torch.backends.cudnn.benchmark = True
working_dir = os.path.dirname(os.path.abspath(__file__))

# ------------------- API Key & Caching Manager -------------------
from llm_manager import get_next_groq_key  # <- NEW

# Select current API key from rotation
groq_api_key = get_next_groq_key(st.session_state)

# ------------------- Lazy Initialization -------------------
@st.cache_resource(show_spinner=False)
def get_easyocr_reader():
    import easyocr
    return easyocr.Reader(["en"], gpu=torch.cuda.is_available())

@st.cache_data(show_spinner=False)
def ensure_nltk():
    import nltk
    nltk.download('wordnet', quiet=True)
    return WordNetLemmatizer()

lemmatizer = ensure_nltk()
reader = get_easyocr_reader()

from courses import COURSES_BY_CATEGORY, RESUME_VIDEOS, INTERVIEW_VIDEOS, get_courses_for_role





FEATURED_COMPANIES = {
    "tech": [
        {
            "name": "Google",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/2/2f/Google_2015_logo.svg",
            "color": "#4285F4",
            "careers_url": "https://careers.google.com",
            "description": "Leading technology company known for search, cloud, and innovation",
            "categories": ["Software", "AI/ML", "Cloud", "Data Science"]
        },
        {
            "name": "Microsoft",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/4/44/Microsoft_logo.svg",
            "color": "#00A4EF",
            "careers_url": "https://careers.microsoft.com",
            "description": "Global leader in software, cloud, and enterprise solutions",
            "categories": ["Software", "Cloud", "Gaming", "Enterprise"]
        },
        {
            "name": "Amazon",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/a/a9/Amazon_logo.svg",
            "color": "#FF9900",
            "careers_url": "https://www.amazon.jobs",
            "description": "E-commerce and cloud computing giant",
            "categories": ["Software", "Operations", "Cloud", "Retail"]
        },
        {
            "name": "Apple",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/f/fa/Apple_logo_black.svg",
            "color": "#555555",
            "careers_url": "https://www.apple.com/careers",
            "description": "Innovation leader in consumer technology",
            "categories": ["Software", "Hardware", "Design", "AI/ML"]
        },
        {
            "name": "Facebook",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/0/05/Facebook_Logo_%282019%29.png",
            "color": "#1877F2",
            "careers_url": "https://www.metacareers.com/",
            "description": "Social media and technology company",
            "categories": ["Software", "Marketing", "Networking", "AI/ML"]
        },
        {
            "name": "Netflix",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/0/08/Netflix_2015_logo.svg",
            "color": "#E50914",
            "careers_url": "https://explore.jobs.netflix.net/careers",
            "description": "Streaming media company",
            "categories": ["Software", "Marketing", "Design", "Service"],
            "website": "https://jobs.netflix.com/",
            "industry": "Entertainment & Technology"
        }
    ],
    "indian_tech": [
        {
            "name": "TCS",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/f/f6/TCS_New_Logo.svg",
            "color": "#0070C0",
            "careers_url": "https://www.tcs.com/careers",
            "description": "India's largest IT services company",
            "categories": ["IT Services", "Consulting", "Digital"]
        },
        {
            "name": "Infosys",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/5/55/Infosys_logo.svg",
            "color": "#007CC3",
            "careers_url": "https://www.infosys.com/careers",
            "description": "Global leader in digital services and consulting",
            "categories": ["IT Services", "Consulting", "Digital"]
        },
        {
            "name": "Wipro",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/8/80/Wipro_Primary_Logo_Color_RGB.svg",
            "color": "#341F65",
            "careers_url": "https://careers.wipro.com",
            "description": "Leading global information technology company",
            "categories": ["IT Services", "Consulting", "Digital"]
        },
        {
            "name": "HCL",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/5/5e/HCL_Technologies_logo.svg",
            "color": "#0075C9",
            "careers_url": "https://www.hcltech.com/careers",
            "description": "Global technology company",
            "categories": ["IT Services", "Engineering", "Digital"]
        }
    ],
    "global_corps": [
        {
            "name": "IBM",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/5/51/IBM_logo.svg",
            "color": "#1F70C1",
            "careers_url": "https://www.ibm.com/careers",
            "description": "Global leader in technology and consulting",
            "categories": ["Software", "Consulting", "AI/ML", "Cloud"],
            "website": "https://www.ibm.com/careers/",
            "industry": "Technology & Consulting"
        },
        {
            "name": "Accenture",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/8/80/Accenture_Logo.svg",
            "color": "#A100FF",
            "careers_url": "https://www.accenture.com/careers",
            "description": "Global professional services company",
            "categories": ["Consulting", "Technology", "Digital"]
        },
        {
            "name": "Cognizant",
            "logo_url": "https://upload.wikimedia.org/wikipedia/commons/6/6e/Cognizant_logo_2022.svg",
            "color": "#1299D8",
            "careers_url": "https://careers.cognizant.com",
            "description": "Leading professional services company",
            "categories": ["IT Services", "Consulting", "Digital"]
        }
    ]
}


JOB_MARKET_INSIGHTS = {
    "trending_skills": [
        {"name": "Artificial Intelligence", "growth": "+45%", "icon": "fas fa-brain"},
        {"name": "Cloud Computing", "growth": "+38%", "icon": "fas fa-cloud"},
        {"name": "Data Science", "growth": "+35%", "icon": "fas fa-chart-line"},
        {"name": "Cybersecurity", "growth": "+32%", "icon": "fas fa-shield-alt"},
        {"name": "DevOps", "growth": "+30%", "icon": "fas fa-code-branch"},
        {"name": "Machine Learning", "growth": "+28%", "icon": "fas fa-robot"},
        {"name": "Blockchain", "growth": "+25%", "icon": "fas fa-lock"},
        {"name": "Big Data", "growth": "+23%", "icon": "fas fa-database"},
        {"name": "Internet of Things", "growth": "+21%", "icon": "fas fa-wifi"}
    ],
    "top_locations": [
        {"name": "Bangalore", "jobs": "50,000+", "icon": "fas fa-city"},
        {"name": "Mumbai", "jobs": "35,000+", "icon": "fas fa-city"},
        {"name": "Delhi NCR", "jobs": "30,000+", "icon": "fas fa-city"},
        {"name": "Hyderabad", "jobs": "25,000+", "icon": "fas fa-city"},
        {"name": "Pune", "jobs": "20,000+", "icon": "fas fa-city"},
        {"name": "Chennai", "jobs": "15,000+", "icon": "fas fa-city"},
        {"name": "Noida", "jobs": "10,000+", "icon": "fas fa-city"},
        {"name": "Vadodara", "jobs": "7,000+", "icon": "fas fa-city"},
        {"name": "Ahmedabad", "jobs": "6,000+", "icon": "fas fa-city"},
        {"name": "Remote", "jobs": "3,000+", "icon": "fas fa-globe-americas"},
    ],
    "salary_insights": [
        {"role": "Machine Learning Engineer", "range": "10-35 LPA", "experience": "0-5 years"},
        {"role": "Big Data Engineer", "range": "8-30 LPA", "experience": "0-5 years"},
        {"role": "Software Engineer", "range": "5-25 LPA", "experience": "0-5 years"},
        {"role": "Data Scientist", "range": "8-30 LPA", "experience": "0-5 years"},
        {"role": "DevOps Engineer", "range": "6-28 LPA", "experience": "0-5 years"},
        {"role": "UI/UX Designer", "range": "5-25 LPA", "experience": "0-5 years"},
        {"role": "Full Stack Developer", "range": "8-30 LPA", "experience": "0-5 years"},
        {"role": "C++/C#/Python/Java Developer", "range": "6-26 LPA", "experience": "0-5 years"},
        {"role": "Django Developer", "range": "7-27 LPA", "experience": "0-5 years"},
        {"role": "Cloud Engineer", "range": "6-26 LPA", "experience": "0-5 years"},
        {"role": "Google Cloud/AWS/Azure Engineer", "range": "6-26 LPA", "experience": "0-5 years"},
        {"role": "Salesforce Engineer", "range": "6-26 LPA", "experience": "0-5 years"},
    ]
}

def get_featured_companies(category=None):
    """Get featured companies with original logos, optionally filtered by category"""
    def has_valid_logo(company):
        return "logo_url" in company and company["logo_url"].startswith("https://upload.wikimedia.org/")

    if category and category in FEATURED_COMPANIES:
        return [company for company in FEATURED_COMPANIES[category] if has_valid_logo(company)]
    
    return [
        company for companies in FEATURED_COMPANIES.values()
        for company in companies if has_valid_logo(company)
    ]


def get_market_insights():
    """Get job market insights"""
    return JOB_MARKET_INSIGHTS

def get_company_info(company_name):
    """Get company information by name"""
    for companies in FEATURED_COMPANIES.values():
        for company in companies:
            if company["name"] == company_name:
                return company
    return None

def get_companies_by_industry(industry):
    """Get list of companies by industry"""
    companies = []
    for companies_list in FEATURED_COMPANIES.values():
        for company in companies_list:
            if "industry" in company and company["industry"] == industry:
                companies.append(company)
    return companies

# Gender-coded language
gender_words = {
    "masculine": [
        "active", "aggressive", "ambitious", "analytical", "assertive", "autonomous", "boast", "bold",
        "challenging", "competitive", "confident", "courageous", "decisive", "determined", "dominant", "driven",
        "dynamic", "forceful", "independent", "individualistic", "intellectual", "lead", "leader", "objective",
        "outspoken", "persistent", "principled", "proactive", "resilient", "self-reliant", "self-sufficient",
        "strong", "superior", "tenacious","guru","tech guru","technical guru", "visionary", "manpower", "strongman", "command",
        "assert", "headstrong", "rockstar", "superstar", "go-getter", "trailblazer", "results-driven",
        "fast-paced", "driven", "determination", "competitive spirit"
    ],
    
    "feminine": [
        "affectionate", "agreeable", "attentive", "collaborative", "committed", "compassionate", "considerate",
        "cooperative", "dependable", "dependent", "emotional", "empathetic", "enthusiastic", "friendly", "gentle",
        "honest", "inclusive", "interpersonal", "kind", "loyal", "modest", "nurturing", "pleasant", "polite",
        "sensitive", "supportive", "sympathetic", "tactful", "tender", "trustworthy", "understanding", "warm",
        "yield", "adaptable", "communal", "helpful", "dedicated", "respectful", "nurture", "sociable",
        "relationship-oriented", "team player", "dependable", "people-oriented", "empathetic listener",
        "gentle communicator", "open-minded"
    ]
}
# Sample job search function
import uuid
import urllib.parse

def search_jobs(job_role, location, experience_level=None, job_type=None, foundit_experience=None):
    # Encode inputs
    role_encoded = urllib.parse.quote_plus(job_role.strip())
    loc_encoded = urllib.parse.quote_plus(location.strip())

    # Mappings
    experience_range_map = {
        "Internship": "0~0", "Entry Level": "1~1", "Associate": "2~3",
        "Mid-Senior Level": "4~7", "Director": "8~15", "Executive": "16~20"
    }

    experience_exact_map = {
        "Internship": "0", "Entry Level": "1", "Associate": "2",
        "Mid-Senior Level": "4", "Director": "8", "Executive": "16"
    }

    linkedin_exp_map = {
        "Internship": "1", "Entry Level": "2", "Associate": "3",
        "Mid-Senior Level": "4", "Director": "5", "Executive": "6"
    }

    job_type_map = {
        "Full-time": "F", "Part-time": "P", "Contract": "C",
        "Temporary": "T", "Volunteer": "V", "Internship": "I"
    }

    # LinkedIn
    linkedin_url = f"https://www.linkedin.com/jobs/search/?keywords={role_encoded}&location={loc_encoded}"
    if experience_level in linkedin_exp_map:
        linkedin_url += f"&f_E={linkedin_exp_map[experience_level]}"
    if job_type in job_type_map:
        linkedin_url += f"&f_JT={job_type_map[job_type]}"

    # Naukri
    # Naukri - add keyword (k), location (l), and experience if available
    naukri_url = f"https://www.naukri.com/{role_encoded}-jobs-in-{loc_encoded}?k={role_encoded}&l={loc_encoded}"
    if experience_level and experience_exact_map.get(experience_level):
     naukri_url += f"&experience={experience_exact_map[experience_level]}"


    # FoundIt
    if foundit_experience is not None:
        experience_range = f"{foundit_experience}~{foundit_experience}"
        experience_exact = str(foundit_experience)
    else:
        experience_range = experience_range_map.get(experience_level, "")
        experience_exact = experience_exact_map.get(experience_level, "")

    search_id = uuid.uuid4()
    foundit_url = f"https://www.foundit.in/srp/results?query={role_encoded}&locations={loc_encoded}"
    if experience_range:
        foundit_url += f"&experienceRanges={urllib.parse.quote_plus(experience_range)}"
    if experience_exact:
        foundit_url += f"&experience={experience_exact}"
    foundit_url += f"&searchId={search_id}"

    return [
        {"title": f"LinkedIn: {job_role} jobs in {location}", "link": linkedin_url},
        {"title": f"Naukri: {job_role} jobs in {location}", "link": naukri_url},
        {"title": f"FoundIt (Monster): {job_role} jobs in {location}", "link": foundit_url}
    ]




def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A function to add a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color and underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    color_element = OxmlElement('w:color')
    color_element.set(qn('w:val'), color)
    rPr.append(color_element)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def generate_docx(text, filename="bias_free_resume.docx"):
    doc = Document()
    doc.add_heading('Bias-Free Resume', 0)
    doc.add_paragraph(text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Extract text from PDF
def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text_list = [page.get_text("text") for page in doc if page.get_text("text").strip()]
        doc.close()
        return text_list if text_list else extract_text_from_images(file_path)
    except Exception as e:
        st.error(f"⚠ Error extracting text: {e}")
        return []

def extract_text_from_images(pdf_path):
    try:
        images = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=5)
        return ["\n".join(reader.readtext(np.array(img), detail=0)) for img in images]
    except Exception as e:
        st.error(f"⚠ Error extracting from image: {e}")
        return []

# Detect bias in resume

def detect_bias(text):
    text = text.lower()
    masc, fem = 0, 0

    masculine_words_sorted = sorted(gender_words["masculine"], key=len, reverse=True)
    feminine_words_sorted = sorted(gender_words["feminine"], key=len, reverse=True)

    for phrase in masculine_words_sorted:
        masc += len(re.findall(rf'\b{re.escape(phrase)}\b', text))

    for phrase in feminine_words_sorted:
        fem += len(re.findall(rf'\b{re.escape(phrase)}\b', text))

    total = masc + fem

    if total == 0:
        return 0.0, masc, fem  

    bias_score = min(total / 20, 1.0)

    return round(bias_score, 2), masc, fem

gender_words = {
    "masculine": [
        "active", "aggressive", "ambitious", "analytical", "assertive", "autonomous", "boast", "bold",
        "challenging", "competitive", "confident", "courageous", "decisive", "determined", "dominant", "driven",
        "dynamic", "forceful", "independent", "individualistic", "intellectual", "lead", "leader", "objective",
        "outspoken", "persistent", "principled", "proactive", "resilient", "self-reliant", "self-sufficient",
        "strong", "superior", "tenacious","guru","tech guru","technical guru", "visionary", "manpower", "strongman", "command",
        "assert", "headstrong", "rockstar", "superstar", "go-getter", "trailblazer", "results-driven",
        "fast-paced", "driven", "determination", "competitive spirit"
    ],
    
    "feminine": [
        "affectionate", "agreeable", "attentive", "collaborative", "committed", "compassionate", "considerate",
        "cooperative", "dependable", "dependent", "emotional", "empathetic", "enthusiastic", "gentle",
        "honest", "inclusive", "interpersonal", "kind", "loyal", "modest", "nurturing", "pleasant", "polite",
        "sensitive", "supportive", "sympathetic", "tactful", "tender", "trustworthy", "understanding", "warm",
        "yield", "adaptable", "communal", "helpful", "dedicated", "respectful", "nurture", "sociable",
        "relationship-oriented", "team player", "dependable", "people-oriented", "empathetic listener",
        "gentle communicator", "open-minded"
    ]
}
replacement_mapping = {
    "masculine": {
        "active": "engaged",
        "aggressive": "proactive",
        "ambitious": "motivated",
        "analytical": "detail-oriented",
        "assertive": "confident",
        "autonomous": "self-directed",
        "boast": "highlight",
        "bold": "confident",
        "challenging": "demanding",
        "competitive": "goal-oriented",
        "confident": "self-assured",
        "courageous": "bold",
        "decisive": "action-oriented",
        "determined": "focused",
        "dominant": "influential",
        "driven": "committed",
        "dynamic": "adaptable",
        "forceful": "persuasive",
        "guru":"technical expert",
        "independent": "self-sufficient",
        "individualistic": "self-motivated",
        "intellectual": "knowledgeable",
        "lead": "guide",
        "leader": "team lead",
        "objective": "unbiased",
        "outspoken": "expressive",
        "persistent": "resilient",
        "principled": "ethical",
        "proactive": "initiative-taking",
        "resilient": "adaptable",
        "self-reliant": "resourceful",
        "self-sufficient": "capable",
        "strong": "capable",
        "superior": "exceptional",
        "tenacious": "determined",
        "technical guru": "technical expert",
        "visionary": "forward-thinking",
        "manpower": "workforce",
        "strongman": "resilient individual",
        "command": "direct",
        "assert": "state confidently",
        "headstrong": "determined",
        "rockstar": "top performer",
        "superstar": "outstanding contributor",
        "go-getter": "initiative-taker",
        "trailblazer": "innovator",
        "results-driven": "outcome-focused",
        "fast-paced": "dynamic",
        "determination": "commitment",
        "competitive spirit": "goal-oriented mindset"
    },
    
    "feminine": {
        "affectionate": "approachable",
        "agreeable": "cooperative",
        "attentive": "observant",
        "collaborative": "team-oriented",
        "collaborate": "team-oriented",
        "collaborated": "worked together",
        "committed": "dedicated",
        "compassionate": "caring",
        "considerate": "thoughtful",
        "cooperative": "supportive",
        "dependable": "reliable",
        "dependent": "team-oriented",
        "emotional": "passionate",
        "empathetic": "understanding",
        "enthusiastic": "positive",
        "gentle": "respectful",
        "honest": "trustworthy",
        "inclusive": "open-minded",
        "interpersonal": "people-focused",
        "kind": "respectful",
        "loyal": "dedicated",
        "modest": "humble",
        "nurturing": "supportive",
        "pleasant": "positive",
        "polite": "professional",
        "sensitive": "attentive",
        "supportive": "encouraging",
        "sympathetic": "understanding",
        "tactful": "diplomatic",
        "tender": "considerate",
        "trustworthy": "reliable",
        "understanding": "empathetic",
        "warm": "welcoming",
        "yield": "adaptable",
        "adaptable": "flexible",
        "communal": "team-centered",
        "helpful": "supportive",
        "dedicated": "committed",
        "respectful": "considerate",
        "nurture": "develop",
        "sociable": "friendly",
        "relationship-oriented": "team-focused",
        "team player": "collaborative member",
        "people-oriented": "person-focused",
        "empathetic listener": "active listener",
        "gentle communicator": "considerate communicator",
        "open-minded": "inclusive"
    }
}

def rewrite_text_with_llm(text, replacement_mapping, user_location):
    """
    Uses LLM to rewrite a resume with bias-free language and suggest relevant job roles.
    Applies strict word replacement mapping and structures the result.
    """
    from llm_manager import call_llm

    # Format the replacement mapping as a readable bullet list for the prompt
    formatted_mapping = "\n".join(
        [f"- \"{key}\" → \"{value}\"" for key, value in replacement_mapping.items()]
    )

    # Construct the prompt
    prompt = f"""
You are an expert career advisor and professional resume language editor.

Your task is to:

1. **Rewrite the following resume text** to:
   - Remove or replace any gender-coded, biased, or non-inclusive language.
   - Use *professional, inclusive, neutral, clear, and grammatically correct language*.
   - **Retain all technical terms, job-specific keywords, certifications, and proper names.**
   - Do **not** add new content or remove important information.
   - Preserve the original meaning and intent of each sentence.

---

2. **Structure and Organize** the rewritten resume into clearly labeled standard resume sections. Only include sections that are present in the original text:
   - Name
   - Contact Information
   - Email
   - Portfolio
   - Professional Summary
   - Work Experience
   - Skills
   - Certifications
   - Education
   - Projects
   - Interests

   - If *Name*, *Contact Information*, or *Email* is present, place them clearly at the top under respective headings.

---

3. **Strictly apply the following word replacement mapping:**

{formatted_mapping}

   - If a word or phrase matches a key exactly from this list, replace it with the corresponding value.
   - Leave all other content unchanged.

---

4. **Suggest 5 suitable job titles** based on the resume content and the candidate’s location: **{user_location}**
   - Ensure titles are realistic for this location and aligned with the candidate's experience and skills.
   - Provide a brief explanation for each suggestion.

---

5. **Provide LinkedIn job search URLs** for each suggested title based on the location: **{user_location}**

---

**Original Resume Text:**
\"\"\"{text}\"\"\"

---

**✅ Bias-Free Rewritten Resume (Well-Structured):**

---

**🎯 Suggested Job Titles with Explanations and LinkedIn URLs:**

1. **Job Title 1** — Reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%201&location={user_location})

2. **Job Title 2** — Reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%202&location={user_location})

3. **Job Title 3** — Reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%203&location={user_location})

4. **Job Title 4** — Reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%204&location={user_location})

5. **Job Title 5** — Reason  
🔗 [Search on LinkedIn](https://www.linkedin.com/jobs/search/?keywords=Job%20Title%205&location={user_location})
"""

    # Call the LLM with caching + key rotation
    response = call_llm(prompt, session=st.session_state)
    return response




def rewrite_and_highlight(text, replacement_mapping, user_location):
    highlighted_text = text
    masculine_count, feminine_count = 0, 0
    detected_masculine_words = Counter()
    detected_feminine_words = Counter()

    words = re.findall(r'\b\w+\b', text)

    for w in words:
        lemma = lemmatizer.lemmatize(w.lower())
        if lemma in gender_words["masculine"]:
            masculine_count += 1
            detected_masculine_words[lemma] += 1
            highlighted_text = re.sub(rf'\b{re.escape(w)}\b', f":blue[{w}]", highlighted_text)
 
        elif lemma in gender_words["feminine"]:
            feminine_count += 1
            detected_feminine_words[lemma] += 1
            highlighted_text = re.sub(rf'\b{re.escape(w)}\b', f":red[{w}]", highlighted_text)

    # Now rewrite the text using the LLM
    rewritten_text = rewrite_text_with_llm(text, replacement_mapping, user_location)



    return highlighted_text, rewritten_text, masculine_count, feminine_count, detected_masculine_words, detected_feminine_words


def ats_percentage_score(
    resume_text,
    job_description,
    logic_profile_score=None,
    edu_weight=20,
    exp_weight=35,
    skills_weight=30,
    lang_weight=5,
    keyword_weight=10
):
    """
    Analyzes resume against job description using a structured, score-based prompt.
    Enforces strict score calculation to prevent inflation. Final scores are validated externally.
    """
    from llm_manager import call_llm

    logic_score_note = (
        f"\n\nOptional Note: The system also calculated a logic-based profile score of {logic_profile_score}/100 based on resume length, experience, and skills."
        if logic_profile_score else ""
    )

    prompt = f"""
You are an AI-powered ATS evaluator. You must evaluate the candidate's resume strictly based on the scoring rules below. 
You are not allowed to guess, assume, or round scores casually. Your component scores will be validated programmatically,
so they must follow **exact arithmetic based on the provided weights**.

---

📊 **Scoring Formula (Total = 100 Points)**

1. **Education Score ({edu_weight} pts)**
   - +50% if degree title (e.g., B.Tech, MSc) matches JD
   - +50% if field (e.g., CS, IT, Engineering) matches JD
   - Partial matches get 25% each

2. **Experience Score ({exp_weight} pts)**
   - +60% if years of experience ≥ required
   - +30% if role titles match
   - +10% if domain/industry matches

3. **Skills Match ({skills_weight} pts)**
   - Extract skill sets from both resume and JD
   - Score = (# of matching skills / total JD skills) × {skills_weight}
   - Show both skill lists

4. **Language Quality ({lang_weight} pts)**
   - {lang_weight} = Very clear, formal, professional tone
   - Half for minor grammatical/style issues
   - Low score for informal or unclear writing

5. **Keyword Match ({keyword_weight} pts)**
   - Extract tools, tech, frameworks from JD
   - For each missing keyword, deduct proportionally from {keyword_weight}
   - Show missing keyword list

---

📐 **Scoring Instruction**
- Total Score = Sum of all 5 components.
- This becomes the **Overall Percentage Match**.
- Max = 100. If score exceeds, cap it at 100.
- Never return 0 unless all components are truly zero.
- Return exact numeric values.

---

🧾 **OUTPUT FORMAT (Strictly follow this):**

Candidate Name: <full name or "Not Found">

Education Score: <0–{edu_weight}>  
Education Match Details: <e.g., "Degree title matches B.Tech; field matches Computer Science.">

Experience Score: <0–{exp_weight}>  
Experience Highlights: <e.g., "5 years experience (JD requires 4), title matches 'Software Engineer', domain matched IT.">

Skills Match Percentage: <0–{skills_weight}>  
Skills Found in Resume: <comma-separated list>  
Skills Required by JD: <comma-separated list>  
Skills Missing: <comma-separated list>

Language Quality Score: <0–{lang_weight}>  
Language Quality Comments: <e.g., "Professional tone with minor grammatical errors.">

Keyword Match Score: <0–{keyword_weight}>  
Missing Keywords: <comma-separated list>

Overall Percentage Match: <sum of above components>  
Formatted Score: <Excellent / Good / Average / Poor>

Final Thoughts:  
Provide a detailed summary (4–6 sentences) about the candidate’s overall fit. Highlight strengths

{logic_score_note}

---

### Job Description:
\"\"\"{job_description}\"\"\"

---

### Resume:
\"\"\"{resume_text}\"\"\"
"""

    response = call_llm(prompt, session=st.session_state)
    return response.strip()


# Setup Vector DB
def setup_vectorstore(documents):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    if DEVICE == "cuda":
        embeddings.model = embeddings.model.to(torch.device("cuda"))
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    doc_chunks = text_splitter.split_text("\n".join(documents))
    return FAISS.from_texts(doc_chunks, embeddings)

# Create Conversational Chain
def create_chain(vectorstore):
    if "memory" not in st.session_state:
        st.session_state.memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

    llm = ChatGroq(model="llama-3.3-70b-versatile", temperature=0, groq_api_key=groq_api_key)
    retriever = vectorstore.as_retriever(search_kwargs={"k": 5})

    return ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        chain_type="stuff",
        memory=st.session_state.memory,
        verbose=False
    )

# App Title
st.title("🦙 Chat with LEXIBOT - LLAMA 3.3 (Bias Detection + QA + GPU)")

# Chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

    st.sidebar.markdown("### 🏷️ Job Information")
job_title = st.sidebar.text_input("💼 Job Title")  # <-- New input for job title

st.sidebar.markdown("### 📝 Paste Job Description")
job_description = st.sidebar.text_area("Enter the Job Description here", height=200)

if job_description.strip() == "":
    st.sidebar.warning("Please enter a job description to evaluate the resumes.")

user_location = st.sidebar.text_input("📍 Preferred Job Location (City, Country)")

st.sidebar.markdown("### 🎛️ Customize ATS Scoring Weights")


edu_weight = st.sidebar.slider("🎓 Education Weight", 0, 50, 20)
exp_weight = st.sidebar.slider("💼 Experience Weight", 0, 50, 35)
skills_weight = st.sidebar.slider("🛠 Skills Match Weight", 0, 50, 30)
lang_weight = st.sidebar.slider("🗣 Language Quality Weight", 0, 10, 5)
keyword_weight = st.sidebar.slider("🔑 Keyword Match Weight", 0, 20, 10)

total_weight = edu_weight + exp_weight + skills_weight + lang_weight + keyword_weight

if total_weight != 100:
    st.sidebar.error(f"⚠️ Total = {total_weight}. Please make it exactly 100.")
else:
    st.sidebar.success("✅ Total weight = 100")


uploaded_files = st.file_uploader("Upload PDF Resumes", type=["pdf"], accept_multiple_files=True)


import os
import re
import streamlit as st
from datetime import datetime
resume_data = []

# 🧠 Dynamic domain detection from title + description
def detect_domain_from_title_and_description(job_title, job_description):
    title = job_title.lower().strip()
    jd = job_description.lower().strip()
    combined = f"{title} {jd}"

    if "full stack" in combined or "fullstack" in combined:
        return "Software Engineering"
    if "frontend" in combined or "react" in combined or "angular" in combined or "vue" in combined:
        return "Frontend"
    if "backend" in combined or "node" in combined or "django" in combined or "api" in combined:
        return "Backend"
    if "software engineer" in combined or "web developer" in combined or "developer" in combined:
        return "Software Engineering"
    if "machine learning" in combined or "ml engineer" in combined or "deep learning" in combined or "ai engineer" in combined:
        return "AI/ML"
    if "data scientist" in combined or "data science" in combined:
        return "Data Science"
    if "cybersecurity" in combined or "security analyst" in combined or "penetration testing" in combined or "owasp" in combined:
        return "Cybersecurity"
    if "cloud" in combined or "aws" in combined or "azure" in combined or "gcp" in combined:
        return "Cloud"
    if "docker" in combined or "kubernetes" in combined or "ci/cd" in combined:
        return "DevOps"
    if "android" in combined or "ios" in combined or "mobile" in combined:
        return "Mobile Development"
    if "ui" in combined or "ux" in combined or "figma" in combined or "designer" in combined or "user interface" in combined:
        return "UI/UX"
    return "General"

# ✏️ Resume Evaluation Logic
if uploaded_files and job_description:
    all_text = []

    for uploaded_file in uploaded_files:
        # Save the uploaded file
        file_path = os.path.join(working_dir, uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        # Extract text from PDF
        text = extract_text_from_pdf(file_path)
        all_text.extend(text)
        full_text = " ".join(text)

        # Bias detection
        bias_score, masc, fem = detect_bias(full_text)
        highlighted_text, rewritten_text, masc_count, fem_count, detected_masc, detected_fem = rewrite_and_highlight(
            full_text, replacement_mapping, user_location
        )

        # ATS scoring
        ats_result = ats_percentage_score(
            resume_text=full_text,
            job_description=job_description,
            logic_profile_score=None,
            edu_weight=edu_weight,
            exp_weight=exp_weight,
            skills_weight=skills_weight,
            lang_weight=lang_weight,
            keyword_weight=keyword_weight
        )

        # Helper extractors
        def extract_score(pattern, text, default=0):
            match = re.search(pattern, text)
            return int(match.group(1)) if match else default

        def extract_text(pattern, text, default="N/A"):
            match = re.search(pattern, text)
            return match.group(1).strip() if match else default

        # Extract info
        candidate_name = extract_text(r"Candidate Name:\s*(.*)", ats_result)
        ats_score = extract_score(r"Overall Percentage Match:\s*(\d+)", ats_result)
        edu_score = extract_score(r"Education Score:\s*(\d+)", ats_result)
        exp_score = extract_score(r"Experience Score:\s*(\d+)", ats_result)
        skills_score = extract_score(r"Skills Match Percentage:\s*(\d+)", ats_result)
        lang_score = extract_score(r"Language Quality Score:\s*(\d+)", ats_result)
        keyword_score = extract_score(r"Keyword Match Score:\s*(\d+)", ats_result)
        formatted_score = extract_text(r"Formatted Score:\s*(.*)", ats_result)
        missing_keywords = extract_text(r"Missing Keywords:\s*(.*)", ats_result)
        fit_summary = extract_text(r"Final Thoughts:\s*(.*)", ats_result)

        # Domain classification
        domain = detect_domain_from_title_and_description(job_title, job_description)

        # Store in dashboard data
        resume_data.append({
            "Resume Name": uploaded_file.name,
            "Candidate Name": candidate_name,
            "ATS Match %": ats_score,
            "Formatted Score": formatted_score,
            "Education Score": edu_score,
            "Experience Score": exp_score,
            "Skills Match %": skills_score,
            "Language Quality Score": lang_score,
            "Keyword Match Score": keyword_score,
            "Missing Keywords": missing_keywords,
            "Fit Summary": fit_summary,
            "Bias Score (0 = Fair, 1 = Biased)": bias_score,
            "Masculine Words": masc_count,
            "Feminine Words": fem_count,
            "Detected Masculine Words": detected_masc,
            "Detected Feminine Words": detected_fem,
            "Text Preview": full_text[:300] + "...",
            "Highlighted Text": highlighted_text,
            "Rewritten Text": rewritten_text
        })

        # ✅ Insert into DB (timestamp is handled by SQLite)
        from db_manager import insert_candidate
        insert_candidate((
            uploaded_file.name,
            candidate_name,
            ats_score,
            edu_score,
            exp_score,
            skills_score,
            lang_score,
            keyword_score,
            bias_score,
            domain
        ))

    # Finalize processing
    st.success("✅ All resumes processed!")

    if all_text:
        st.session_state.vectorstore = setup_vectorstore(all_text)
        st.session_state.chain = create_chain(st.session_state.vectorstore)

# === TAB 1: Dashboard ===
# 📊 Dashboard and Metrics
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📊 Dashboard", "🧾 Resume Builder", "💼 Job Search", 
    "📚 Course Recommendation", "📁 Admin DB View"
])


# === TAB 1: Dashboard ===
with tab1:
    if resume_data:
        total_masc = sum(r["Masculine Words"] for r in resume_data)
        total_fem = sum(r["Feminine Words"] for r in resume_data)
        avg_bias = round(np.mean([r["Bias Score (0 = Fair, 1 = Biased)"] for r in resume_data]), 2)
        total_resumes = len(resume_data)

        st.markdown("### 📊 Summary Statistics")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("📄 Resumes Uploaded", total_resumes)
        with col2:
            st.metric("🔎 Avg. Bias Score", avg_bias)
        with col3:
            st.metric("🔵 Total Masculine Words", total_masc)
        with col4:
            st.metric("🔴 Total Feminine Words", total_fem)

        st.markdown("### 🗂️ Resumes Overview")
        df = pd.DataFrame(resume_data)
        st.dataframe(
            df[[ 
                "Resume Name", "Candidate Name", "ATS Match %", "Education Score",
                "Experience Score", "Skills Match %", "Language Quality Score", "Keyword Match Score",
                "Bias Score (0 = Fair, 1 = Biased)", "Masculine Words", "Feminine Words"
            ]],
            use_container_width=True
        )

        st.markdown("### 📊 Visual Analysis")
        chart_tab1, chart_tab2 = st.tabs(["📉 Bias Score Chart", "⚖ Gender-Coded Words"])

        with chart_tab1:
            st.subheader("Bias Score Comparison Across Resumes")
            st.bar_chart(df.set_index("Resume Name")[["Bias Score (0 = Fair, 1 = Biased)"]])

        with chart_tab2:
            st.subheader("Masculine vs Feminine Word Usage")
            fig, ax = plt.subplots(figsize=(10, 5))
            index = np.arange(len(df))
            bar_width = 0.35

            ax.bar(index, df["Masculine Words"], bar_width, label="Masculine", color="#3498db")
            ax.bar(index + bar_width, df["Feminine Words"], bar_width, label="Feminine", color="#e74c3c")

            ax.set_xlabel("Resumes", fontsize=12)
            ax.set_ylabel("Word Count", fontsize=12)
            ax.set_title("Gender-Coded Word Usage per Resume", fontsize=14)
            ax.set_xticks(index + bar_width / 2)
            ax.set_xticklabels(df["Resume Name"], rotation=45, ha='right')
            ax.legend()
            st.pyplot(fig)

        st.markdown("### 📝 Detailed Resume Reports")
        for resume in resume_data:
            with st.expander(f"📄 {resume['Resume Name']} | {resume['Candidate Name']}"):
                st.markdown("### 📊 ATS Evaluation for: **" + resume['Candidate Name'] + "**")
                score_col1, score_col2, score_col3 = st.columns(3)
                with score_col1:
                    st.metric("📈 Overall Match", f"{resume['ATS Match %']}%")
                with score_col2:
                    st.metric("🏆 Formatted Score", resume.get("Formatted Score", "N/A"))
                with score_col3:
                    st.metric("🧠 Language Quality", f"{resume.get('Language Quality Score', 'N/A')} / {lang_weight}")

                col_a, col_b, col_c, col_d = st.columns(4)
                with col_a:
                    st.metric("🎓 Education Score", f"{resume.get('Education Score', 'N/A')} / {edu_weight}")
                with col_b:
                    st.metric("💼 Experience Score", f"{resume.get('Experience Score', 'N/A')} / {exp_weight}")
                with col_c:
                    st.metric("🛠 Skills Match", f"{resume.get('Skills Match %', 'N/A')} / {skills_weight}")
                with col_d:
                    st.metric("🔍 Keyword Score", f"{resume.get('Keyword Match Score', 'N/A')} / {keyword_weight}")

                st.markdown("**❗ Missing Keywords:**")
                missing_list = resume["Missing Keywords"].split(",") if resume["Missing Keywords"] else []
                if missing_list and any(kw.strip() for kw in missing_list):
                    for kw in missing_list:
                        st.error(f"- {kw.strip()}")
                else:
                    st.info("No missing keywords detected.")

                st.markdown("### 📝 Fit Summary")
                st.write(resume['Fit Summary'])

                st.divider()

                detail_tab1, detail_tab2 = st.tabs(["🔎 Bias Analysis", "✅ Rewritten Resume"])

                with detail_tab1:
                    st.markdown("#### Bias-Highlighted Original Text")
                    st.markdown(resume["Highlighted Text"], unsafe_allow_html=True)

                    st.markdown("### 📌 Gender-Coded Word Counts:")
                    bias_col1, bias_col2 = st.columns(2)
                    with bias_col1:
                        st.metric("🔵 Masculine Words", resume["Masculine Words"])
                        if resume["Detected Masculine Words"]:
                            st.markdown("### 📚 Detected Words:")
                            st.success(", ".join(f"{word} ({count})" for word, count in resume["Detected Masculine Words"].items()))
                        else:
                            st.info("No masculine words detected.")
                    with bias_col2:
                        st.metric("🔴 Feminine Words", resume["Feminine Words"])
                        if resume["Detected Feminine Words"]:
                            st.markdown("### 📚 Detected Words:")
                            st.success(", ".join(f"{word} ({count})" for word, count in resume["Detected Feminine Words"].items()))
                        else:
                            st.info("No feminine words detected.")

                with detail_tab2:
                    st.markdown("#### ✨ Bias-Free Rewritten Resume")
                    st.write(resume["Rewritten Text"])

                    docx_file = generate_docx(resume["Rewritten Text"])
                    st.download_button(
    label="📥 Download Bias-Free Resume (.docx)",
    data=docx_file,
    file_name=f"{resume['Resume Name'].split('.')[0]}_bias_free.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True,
    key=f"download_docx_{resume['Resume Name']}"  # ✅ Ensure uniqueness
)


                def generate_resume_report_html(resume):
                    rewritten_text = resume['Rewritten Text'].replace("\n", "<br>")
                    masculine_words = ", ".join(f"{k}({v})" for k, v in resume['Detected Masculine Words'].items())
                    feminine_words = ", ".join(f"{k}({v})" for k, v in resume['Detected Feminine Words'].items())
                    missing_keywords = "".join(
                        f"<span class='keyword'>{kw.strip()}</span>"
                        for kw in resume['Missing Keywords'].split(",") if kw.strip()
                    ) or "<i>None</i>"

                    return f"""
                    <!DOCTYPE html>
                    <html lang="en">
                    <head>
                        <meta charset="UTF-8">
                        <title>{resume['Candidate Name']} - Resume Analysis Report</title>
                        <style>
                            body {{
                                font-family: 'Segoe UI', sans-serif;
                                margin: 40px;
                                background-color: #f5f7fa;
                                color: #333;
                            }}
                            h1, h2 {{
                                color: #2f4f6f;
                            }}
                            .section {{
                                margin-bottom: 30px;
                            }}
                            .highlight {{
                                background-color: #eef;
                                padding: 10px;
                                border-radius: 6px;
                                margin-top: 10px;
                                font-size: 14px;
                            }}
                            .metric-box {{
                                display: inline-block;
                                background: #dbeaff;
                                padding: 10px 20px;
                                margin: 10px;
                                border-radius: 10px;
                                font-weight: bold;
                            }}
                            .keyword {{
                                display: inline-block;
                                background: #fbdcdc;
                                color: #a33;
                                margin: 4px;
                                padding: 6px 12px;
                                border-radius: 12px;
                                font-size: 13px;
                            }}
                            .resume-box {{
                                background-color: #f9f9ff;
                                padding: 15px;
                                border-radius: 8px;
                                border: 1px solid #ccc;
                                white-space: pre-wrap;
                            }}
                        </style>
                    </head>
                    <body>

                        <h1>📄 Resume Analysis Report</h1>

                        <div class="section">
                            <h2>Candidate: {resume['Candidate Name']}</h2>
                            <p><strong>Resume File:</strong> {resume['Resume Name']}</p>
                        </div>

                        <div class="section">
                            <h2>📊 ATS Evaluation</h2>
                            <div class="metric-box">ATS Match: {resume['ATS Match %']}%</div>
                            <div class="metric-box">Education: {resume['Education Score']}</div>
                            <div class="metric-box">Experience: {resume['Experience Score']}</div>
                            <div class="metric-box">Skills Match: {resume['Skills Match %']}</div>
                            <div class="metric-box">Language Score: {resume['Language Quality Score']}</div>
                            <div class="metric-box">Keyword Score: {resume['Keyword Match Score']}</div>
                        </div>

                        <div class="section">
                            <h2>⚖️ Gender Bias Analysis</h2>
                            <div class="metric-box" style="background:#f0f8ff;">Masculine Words: {resume['Masculine Words']}</div>
                            <div class="metric-box" style="background:#fff0f5;">Feminine Words: {resume['Feminine Words']}</div>
                            <p><strong>Bias Score (0=Fair, 1=Biased):</strong> {resume['Bias Score (0 = Fair, 1 = Biased)']}</p>
                            <div class="highlight"><strong>Masculine Words:</strong><br>{masculine_words}</div>
                            <div class="highlight"><strong>Feminine Words:</strong><br>{feminine_words}</div>
                        </div>

                        <div class="section">
                            <h2>📌 Missing Keywords</h2>
                            {missing_keywords}
                        </div>

                        <div class="section">
                            <h2>🧠 Final Fit Summary</h2>
                            <div class="resume-box">{resume['Fit Summary']}</div>
                        </div>

                        <div class="section">
                            <h2>✅ Rewritten Bias-Free Resume</h2>
                            <div class="resume-box">{rewritten_text}</div>
                        </div>

                    </body>
                    </html>
                    """

                html_report = generate_resume_report_html(resume)
                st.download_button(
                    label="📥 Download Full Analysis Report (.html)",
                    data=html_report,
                    file_name=f"{resume['Resume Name'].split('.')[0]}_report.html",
                    mime="text/html",
                    use_container_width=True,
                )
    else:
        st.warning("⚠️ Please upload resumes to view dashboard analytics.")



   
with tab2:
    st.markdown("## 🧾 <span style='color:#336699;'>Advanced Resume Builder</span>", unsafe_allow_html=True)
    st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)
    uploaded_image = st.file_uploader("Upload a Profile Image", type=["png", "jpg", "jpeg"])

    # Initialize session state
    fields = ["name", "email", "phone", "linkedin", "location", "portfolio", "summary", "skills", "languages", "interests","Softskills"]
    for f in fields:
        st.session_state.setdefault(f, "")
    st.session_state.setdefault("experience_entries", [{"title": "", "company": "", "duration": "", "description": ""}])
    st.session_state.setdefault("education_entries", [{"degree": "", "institution": "", "year": "", "details": ""}])
    st.session_state.setdefault("project_entries", [{"title": "", "tech": "", "duration": "", "description": ""}])
    st.session_state.setdefault("project_links", [])
    st.session_state.setdefault("certificate_links", [{"name": "", "link": "", "duration": "", "description": ""}])
    
    # Sidebar controls
    with st.sidebar:
        st.markdown("### ➕ Add More Sections")
        if st.button("➕ Add Experience"):
            st.session_state.experience_entries.append({"title": "", "company": "", "duration": "", "description": ""})
        if st.button("➕ Add Education"):
            st.session_state.education_entries.append({"degree": "", "institution": "", "year": "", "details": ""})
        if st.button("➕ Add Project"):
            st.session_state.project_entries.append({"title": "", "tech": "", "duration": "", "description": ""})
        if st.button("➕ Add Certificate"):
           st.session_state.certificate_links.append({"name": "", "link": "", "duration": "", "description": ""})


    with st.form("resume_form"):
        st.markdown("### 👤 <u>Personal Information</u>", unsafe_allow_html=True)
        with st.container():
            col1, col2 = st.columns(2)
            with col1:
                st.text_input("👤 Full Name ", key="name")
                st.text_input("📞 Phone Number", key="phone")
                st.text_input("📍 Location", key="location")
            with col2:
                st.text_input("📧 Email", key="email")
                st.text_input("🔗 LinkedIn", key="linkedin")
                st.text_input("🌐 Portfolio", key="portfolio")
                st.text_input("💼 Job Title", key="job_title")


        st.markdown("### 📝 <u>Professional Summary</u>", unsafe_allow_html=True)
        st.text_area("Summary", key="summary")

        st.markdown("### 💼 <u>Skills, Languages, Interests & Soft Skills</u>", unsafe_allow_html=True)
        st.text_area("Skills (comma-separated)", key="skills")
        st.text_area("Languages (comma-separated)", key="languages")
        st.text_area("Interests (comma-separated)", key="interests")
        st.text_area("Softskills (comma-separated)", key="Softskills")


        st.markdown("### 🧱 <u>Work Experience</u>", unsafe_allow_html=True)
        for idx, exp in enumerate(st.session_state.experience_entries):
            with st.expander(f"Experience #{idx+1}", expanded=True):
                exp["title"] = st.text_input(f"Job Title", value=exp["title"], key=f"title_{idx}")
                exp["company"] = st.text_input(f"Company", value=exp["company"], key=f"company_{idx}")
                exp["duration"] = st.text_input(f"Duration", value=exp["duration"], key=f"duration_{idx}")
                exp["description"] = st.text_area(f"Description", value=exp["description"], key=f"description_{idx}")

        st.markdown("### 🎓 <u>Education</u>", unsafe_allow_html=True)
        for idx, edu in enumerate(st.session_state.education_entries):
            with st.expander(f"Education #{idx+1}", expanded=True):
                edu["degree"] = st.text_input(f"Degree", value=edu["degree"], key=f"degree_{idx}")
                edu["institution"] = st.text_input(f"Institution", value=edu["institution"], key=f"institution_{idx}")
                edu["year"] = st.text_input(f"Year", value=edu["year"], key=f"edu_year_{idx}")
                edu["details"] = st.text_area(f"Details", value=edu["details"], key=f"edu_details_{idx}")

        st.markdown("### 🛠 <u>Projects</u>", unsafe_allow_html=True)
        for idx, proj in enumerate(st.session_state.project_entries):
            with st.expander(f"Project #{idx+1}", expanded=True):
                proj["title"] = st.text_input(f"Project Title", value=proj["title"], key=f"proj_title_{idx}")
                proj["tech"] = st.text_input(f"Tech Stack", value=proj["tech"], key=f"proj_tech_{idx}")
                proj["duration"] = st.text_input(f"Duration", value=proj["duration"], key=f"proj_duration_{idx}")
                proj["description"] = st.text_area(f"Description", value=proj["description"], key=f"proj_desc_{idx}")


        st.markdown("### 🔗 Project Links")
        project_links_input = st.text_area("Enter one project link per line:")
        if project_links_input:
            st.session_state.project_links = [link.strip() for link in project_links_input.splitlines() if link.strip()]

        st.markdown("### 🧾 <u>Certificates</u>", unsafe_allow_html=True)
        for idx, cert in enumerate(st.session_state.certificate_links):
            with st.expander(f"Certificate #{idx+1}", expanded=True):
                cert["name"] = st.text_input(f"Certificate Name", value=cert["name"], key=f"cert_name_{idx}")
                cert["link"] = st.text_input(f"Certificate Link", value=cert["link"], key=f"cert_link_{idx}")
                cert["duration"] = st.text_input(f"Duration", value=cert["duration"], key=f"cert_duration_{idx}")
                cert["description"] = st.text_area(f"Description", value=cert["description"], key=f"cert_description_{idx}")


        submitted = st.form_submit_button("📑 Generate Resume")

    
        if submitted:
         st.success("✅ Resume Generated Successfully! Scroll down to preview or download.")

        st.markdown("""
    <style>
        .heading-large {
            font-size: 36px;
            font-weight: bold;
            color: #336699;
        }
        .subheading-large {
            font-size: 30px;
            font-weight: bold;
            color: #336699;
        }
        .tab-section {
            margin-top: 20px;
        }
    </style>
""", unsafe_allow_html=True)
 


    # --- Visual Resume Preview Section ---
        st.markdown("## 🧾 <span style='color:#336699;'>Resume Preview</span>", unsafe_allow_html=True)
        st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

        left, right = st.columns([1, 2])

        with left:
            st.markdown(f"""
                <h2 style='color:#2f2f2f;margin-bottom:0;'>{st.session_state['name']}</h2>
                <h4 style='margin-top:5px;color:#444;'>{st.session_state['job_title']}</h4>

                <p style='font-size:14px;'>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}" target="_blank">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}" target="_blank">Portfolio</a>
                </p>
            """, unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Skills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for skill in [s.strip() for s in st.session_state["skills"].split(",") if s.strip()]:
                st.markdown(f"<div style='margin-left:10px;'>• {skill}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Languages</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for lang in [l.strip() for l in st.session_state["languages"].split(",") if l.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {lang}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Interests</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for interest in [i.strip() for i in st.session_state["interests"].split(",") if i.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {interest}</div>", unsafe_allow_html=True)

            st.markdown("<h4 style='color:#336699;'>Softskills</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for Softskills  in [i.strip() for i in st.session_state["Softskills"].split(",") if i.strip()]:
               st.markdown(f"<div style='margin-left:10px;'>• {Softskills}</div>", unsafe_allow_html=True)   


        with right:
            st.markdown("<h4 style='color:#336699;'>Summary</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            summary_text = st.session_state['summary'].replace('\n', '<br>')
            st.markdown(f"<p style='font-size:14px;'>{summary_text}</p>", unsafe_allow_html=True)


            st.markdown("<h4 style='color:#336699;'>Experience</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for exp in st.session_state.experience_entries:
             if exp["company"] or exp["title"]:
              st.markdown(f"""
            <div style='margin-bottom:15px; padding:10px; border-radius:8px;'>
                <div style='display:flex; justify-content:space-between;'>
                    <b>🏢 {exp['company']}</b><span style='color:gray;'>📆  {exp['duration']}</span>
                </div>
                <div style='font-size:14px;'>💼 <i>{exp['title']}</i></div>
                <div style='font-size:14px;'>📝 {exp['description']}</div>
            </div>
        """, unsafe_allow_html=True)


            st.markdown("<h4 style='color:#336699;'>🎓 Education</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for edu in st.session_state.education_entries:
             if edu["institution"] or edu["degree"]:
              st.markdown(f"""
            <div style='margin-bottom: 15px; padding: 10px 15px;color: white; border-radius: 8px;'>
                <div style='display: flex; justify-content: space-between; font-size: 16px; font-weight: bold;'>
                    <span>🏫 {edu['institution']}</span>
                    <span style='color: gray;'>📅 {edu['year']}</span>
                </div>
                <div style='font-size: 14px; margin-top: 5px;'>🎓 <i>{edu['degree']}</i></div>
                <div style='font-size: 14px;'>📄 {edu['details']}</div>
            </div>
        """, unsafe_allow_html=True)


            st.markdown("<h4 style='color:#336699;'>Projects</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
            for proj in st.session_state.project_entries:
             st.markdown(f"""
        <div style='margin-bottom:15px; padding: 10px;'>
        <strong style='font-size:16px;'>{proj['title']}</strong><br>
        <span style='font-size:14px;'>🛠️ <strong>Tech Stack:</strong> {proj['tech']}</span><br>
        <span style='font-size:14px;'>⏳ <strong>Duration:</strong> {proj['duration']}</span><br>
        <span style='font-size:14px;'>📝 <strong>Description:</strong> {proj['description']}</span>
    </div>
    """, unsafe_allow_html=True)



        if st.session_state.project_links:
                st.markdown("<h4 style='color:#336699;'>Project Links</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                for i, link in enumerate(st.session_state.project_links):
                    st.markdown(f"[🔗 Project {i+1}]({link})", unsafe_allow_html=True)

        if st.session_state.certificate_links:
                st.markdown("<h4 style='color:#336699;'>Certificates</h4><hr style='margin-top:-10px;'>", unsafe_allow_html=True)
                
                for cert in st.session_state.certificate_links:
                    if cert["name"] and cert["link"]:
                      st.markdown(f"""
            <div style='display:flex; justify-content:space-between;'>
                <a href="{cert['link']}" target="_blank"><b>📄 {cert['name']}</b></a><span style='color:gray;'>{cert['duration']}</span>
            </div>
            <div style='margin-bottom:10px; font-size:14px;'>{cert['description']}</div>
        """, unsafe_allow_html=True)

# SKILLS
skills_html = "".join(
    f"""
    <div style='display:inline-block; background-color:#e6f0fa; color:#333; 
                padding:8px 16px; margin:6px 6px 6px 0; 
                border-radius:20px; font-size:15px; font-weight:500;'>
        {s.strip()}
    </div>
    """
    for s in st.session_state['skills'].split(',')
    if s.strip()
)
import streamlit as st
import base64

# Initialize profile image HTML with empty string
profile_img_html = ""

if uploaded_image:
    encoded_image = base64.b64encode(uploaded_image.read()).decode()

    # Circular profile image with glowing effect
    profile_img_html = f"""
    <div style="display: flex; justify-content: flex-end; margin-top: 20px;">
        <img src="data:image/png;base64,{encoded_image}" alt="Profile Photo"
             style="
                width: 120px;
                height: 120px;
                border-radius: 50%;
                object-fit: cover;
                border: 3px solid #4da6ff;
                box-shadow:
                    0 0 8px #4da6ff,
                    0 0 16px #4da6ff,
                    0 0 24px #4da6ff;
            " />
    </div>
    """
    st.markdown(profile_img_html, unsafe_allow_html=True)
else:
    st.info("Please upload a profile photo.")



languages_html = "".join(
    f"""
    <div style='display:inline-block; background-color:#e6f0fa; color:#333; 
                padding:8px 16px; margin:6px 6px 6px 0; 
                border-radius:20px; font-size:15px; font-weight:500;'>
        {lang.strip()}
    </div>
    """
    for lang in st.session_state['languages'].split(',')
    if lang.strip()
)


# INTERESTS
interests_html = "".join(
    f"""
    <div style='display:inline-block; background-color:#e6f0fa; color:#333; 
                padding:8px 16px; margin:6px 6px 6px 0; 
                border-radius:20px; font-size:15px; font-weight:500;'>
        {interest.strip()}
    </div>
    """
    for interest in st.session_state['interests'].split(',')
    if interest.strip()
)


Softskills_html = "".join(
    f"""
    <div style='display:inline-block; background-color:#eef3f8; color:#1a1a1a; 
                padding:8px 18px; margin:6px 6px 6px 0; 
                border-radius:25px; font-size:14.5px; font-family:"Segoe UI", sans-serif; 
                font-weight:500; box-shadow: 1px 1px 3px rgba(0,0,0,0.05);'>
        {skill.strip().capitalize()}
    </div>
    """
    for skill in st.session_state['Softskills'].split(',')
    if skill.strip()
)



# EXPERIENCE
experience_html = ""
for exp in st.session_state.experience_entries:
    if exp["company"] or exp["title"]:
        # Handle paragraphs and single line breaks
        description_lines = [line.strip() for line in exp["description"].strip().split("\n\n")]
        description_html = "".join(
            f"<div style='margin-bottom: 8px;'>{line.replace(chr(10), '<br>')}</div>"
            for line in description_lines if line
        )

        experience_html += f"""
        <div style='
            margin-bottom: 20px;
            padding: 16px 20px;
            border-radius: 12px;
            background-color: #dbeaff;
            box-shadow: 0 3px 8px rgba(0, 0, 0, 0.05);
            font-family: "Segoe UI", Tahoma, Geneva, Verdana, sans-serif;
            color: #0a1a33;
            line-height: 1.35;
        '>
            <!-- Header Shadow Card -->
            <div style='
                background-color: #e6f0ff;
                border-radius: 8px;
                padding: 10px 14px;
                margin-bottom: 12px;
                box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            '>
                <div style='
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    font-weight: 600;
                    font-size: 16.5px;
                    margin-bottom: 6px;
                    color: #08244c;
                '>
                    <span>🏢 {exp['company']}</span>
                    <span style='color: #1a2d4f; font-size: 14px;'>📆 {exp['duration']}</span>
                </div>

                <div style='
                    font-size: 16px;
                    font-weight: 700;
                    color: #0b2545;
                '>
                    💼 {exp['title']}
                </div>
            </div>

            <!-- Description -->
            <div style='
                font-size: 15px;
                font-weight: 500;
                color: #102a43;
                line-height: 1.35;
            '>
                📝 {description_html}
            </div>
        </div>
        """


# Convert experience to list if multiple lines

# Escape HTML and convert line breaks
summary_html = st.session_state['summary'].replace('\n', '<br>')


# EDUCATION
education_html = ""
for edu in st.session_state.education_entries:
    if edu.get("institution") or edu.get("details"):
        degree_text = ""
        if edu.get("degree"):
            degree_val = edu["degree"]
            if isinstance(degree_val, list):
                degree_val = ", ".join(degree_val)
            degree_text = f"<div style='font-size: 14px; color: #273c75; margin-bottom: 6px;'>🎓 <b>{degree_val}</b></div>"

        education_html += f"""
        <div style='
            margin-bottom: 20px;
            padding: 16px 20px;
            border-radius: 12px;
            background-color: #e3ebf8;  /* Light Gray Blue */
            box-shadow: 0 3px 8px rgba(39, 60, 117, 0.15);
            font-family: "Segoe UI", Tahoma, Geneva, Verdana, sans-serif;
            color: #273c75;  /* Dark Blue */
            line-height: 1.4;
        '>
            <div style='
                display: flex;
                justify-content: space-between;
                align-items: center;
                font-size: 16px;
                font-weight: 700;
                margin-bottom: 8px;
            '>
                <span>🏫 {edu.get('institution', '')}</span>
                <span style='font-weight: 500;'>🗓️  {edu.get('year', '')}</span>
            </div>
            {degree_text}
            <div style='font-size: 14px; font-style: italic;'>
                📝 {edu.get('details', '')}
            </div>
        </div>
        """




# PROJECTS
# PROJECTS
projects_html = ""
for proj in st.session_state.project_entries:
    if proj.get("title") or proj.get("description"):
        tech_val = proj.get("tech")
        if isinstance(tech_val, list):
            tech_val = ", ".join(tech_val)
        tech_text = f"<div style='font-size: 14px; color: #1b2330; margin-bottom: 8px; text-shadow: 1px 1px 2px rgba(0,0,0,0.15);'><b>🛠️ Technologies:</b> {tech_val if tech_val else ''}</div>" if tech_val else ""

        description_items = ""
        if proj.get("description"):
            description_lines = [line.strip() for line in proj["description"].splitlines() if line.strip()]
            description_items = "".join(f"<li>{line}</li>" for line in description_lines)

        projects_html += f"""
        <div style='
            margin-bottom: 22px;
            padding: 18px 24px;
            border-radius: 14px;
            background-color: #d7e1ec;  /* Deep Blue Slate */
            box-shadow: 0 4px 10px rgba(27, 35, 48, 0.15);
            font-family: "Roboto", "Helvetica Neue", Arial, sans-serif;
            color: #1b2330;  /* Deep Slate Blue */
            line-height: 1.5;
        '>
            <div style='
                font-size: 17px;
                font-weight: 700;
                margin-bottom: 10px;
                display: flex;
                justify-content: space-between;
                align-items: center;
                color: #141a22;
                text-shadow: 1px 1px 2px rgba(0,0,0,0.15);
            '>
                <span>💻 {proj.get('title', '')}</span>
                <span style='font-weight: 600; font-size: 14.5px; text-shadow: 1px 1px 2px rgba(0,0,0,0.15);'>⏳ {proj.get('duration', '')}</span>
            </div>
            {tech_text}
            <div style='font-size: 15px; color: #1b2330; text-shadow: 1px 1px 2px rgba(0,0,0,0.15);'>
                <b>📝 Description:</b>
                <ul style='margin-top: 6px; padding-left: 22px; color: #1b2330;'>
                    {description_items}
                </ul>
            </div>
        </div>
        """




# PROJECT LINKS
project_links_html = ""
if st.session_state.project_links:
    project_links_html = "<h4 class='section-title'>Project Links</h4><hr>" + "".join(
        f'<p><a href="{link}">🔗 Project {i+1}</a></p>'
        for i, link in enumerate(st.session_state.project_links)
    )



certificate_links_html = ""
if st.session_state.certificate_links:
    certificate_links_html = "<h4 class='section-title'>Certificates</h4><hr>"
    for cert in st.session_state.certificate_links:
        if cert["name"] and cert["link"]:
            description = cert.get('description', '').replace('\n', '<br>')
            name = cert['name']
            link = cert['link']
            duration = cert.get('duration', '')

            card_html = f"""
            <div style='
                background-color: #f9fbe7;  /* Green-Yellow pastel */
                padding: 20px 24px;
                border-radius: 16px;
                margin-bottom: 22px;
                box-shadow: 0 4px 14px rgba(34, 60, 80, 0.15);
                font-family: "Poppins", "Segoe UI", sans-serif;
                color: #1a1a1a;
                position: relative;
                line-height: 1.6;
            '>
                <!-- Duration Top Right -->
                <div style='
                    position: absolute;
                    top: 18px;
                    right: 24px;
                    font-size: 13.5px;
                    font-weight: 600;
                    color: #37474f;
                    text-shadow: 0.5px 0.5px 1px rgba(0, 0, 0, 0.15);

                    /* Added background and shadow */
                    background-color: #fffde7;  /* pastel yellow */
                    padding: 4px 12px;
                    border-radius: 14px;
                    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.12);
                '>⏳ {duration}</div>

                <!-- Certificate Title -->
                <div style='
                    font-size: 17px;
                    font-weight: 700;
                    color: #263238;
                    margin-bottom: 8px;
                    text-shadow: 0.5px 0.5px 1.5px rgba(0, 0, 0, 0.1);
                '>
                    📄 <a href="{link}" target="_blank" style='
                        color: #263238;
                        text-decoration: none;
                    '>{name}</a>
                </div>

                <!-- Description -->
                <div style='
                    font-size: 15px;
                    color: #37474f;
                    margin-top: 6px;
                    text-shadow: 0 0 1px rgba(0, 0, 0, 0.08);
                '>
                    📝 {description}
                </div>
            </div>
            """
            certificate_links_html += card_html




        # --- Word Export Logic (Unchanged from your code) ---
html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{st.session_state['name']} - Resume</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            margin: 40px;
            color: #2f2f2f;
        }}
        h2 {{
            font-size: 32px;
            margin: 0;
            color: #336699;
        }}
        h4 {{
            font-size: 24px;
            margin: 0;
            color: #336699;
        }}
        a {{
            color: #007acc;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        hr {{
            border: none;
            border-top: 2px solid #bbb;
            margin: 20px 0;
        }}
        .container {{
            display: flex;
            gap: 40px;
        }}
        .left {{
            flex: 1;
            border-right: 2px solid #ccc;
            padding-right: 20px;
        }}
        .right {{
            flex: 2;
            padding-left: 20px;
        }}
        .section-title {{
            color: #336699;
            margin-top: 30px;
            margin-bottom: 5px;
        }}
        .skill-list {{
            margin-left: 10px;
        }}
        .entry {{
            margin-bottom: 15px;
        }}
        .entry-header {{
            display: flex;
            justify-content: space-between;
        }}
        .entry-title {{
            font-style: italic;
        }}
    </style>
</head>
<body>

    <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px;">
    <div>
        <h2 style="margin: 0; color: #336699;">{st.session_state['name']}</h2>
        <h4 style="margin: 5px 0 0 0; color: #336699;">{st.session_state['job_title']}</h4>
    </div>
    <div>
        {profile_img_html}
    </div>
</div>
<hr>


    <div class="container">
        <div class="left">
            <p>
                📍 {st.session_state['location']}<br>
                📞 {st.session_state['phone']}<br>
                📧 <a href="mailto:{st.session_state['email']}">{st.session_state['email']}</a><br>
                🔗 <a href="{st.session_state['linkedin']}">LinkedIn</a><br>
                🌐 <a href="{st.session_state['portfolio']}">Portfolio</a>
            </p>

            <h4 class="section-title">Skills</h4>
            <hr>
            {skills_html}

            <h4 class="section-title">Languages</h4>
            <hr>
            {languages_html}

            <h4 class="section-title">Interests</h4>
            <hr>
            {interests_html}

            <h4 class="section-title">Softskills</h4>
            <hr>
            {Softskills_html}
        </div>

        <div class="right">
            <h4 class="section-title">Summary</h4>
            <hr>
            <p>{summary_html}</p>

            <h4 class="section-title">Experience</h4>
            <hr>
            {experience_html}

            <h4 class="section-title">Education</h4>
            <hr>
            {education_html}

            <h4 class="section-title">Projects</h4>
            <hr>
            {projects_html}

            {project_links_html}
            {certificate_links_html}
        </div>
    </div>
</body>
</html>
"""

# Then encode it to bytes and prepare for download
html_bytes = html_content.encode("utf-8")
html_file = BytesIO(html_bytes)


with tab2:
 st.download_button (
    label="📥 Download Resume (HTML)",
    data=html_file,
    file_name=f"{st.session_state['name'].replace(' ', '_')}_Resume.html",
    mime="text/html"
)    
with tab2:
 st.markdown("""
✅ After downloading your HTML resume, you can [click here to convert it to PDF](https://www.sejda.com/html-to-pdf) using Sejda's free online tool.
""")

with tab3:
    st.header("🔍 Job Search Across LinkedIn, Naukri, and FoundIt")

    col1, col2 = st.columns(2)

    with col1:
        job_role = st.text_input("💼 Desired Job Role", placeholder="e.g., Data Scientist")
        experience_level = st.selectbox(
            "📈 Experience Level",
            ["", "Internship", "Entry Level", "Associate", "Mid-Senior Level", "Director", "Executive"]
        )

    with col2:
        location = st.text_input("📍 Preferred Location", placeholder="e.g., Bangalore, India")
        job_type = st.selectbox(
            "📋 Job Type",
            ["", "Full-time", "Part-time", "Contract", "Temporary", "Volunteer", "Internship"]
        )

    foundit_experience = st.text_input("🔢 Experience (Years) for FoundIt", placeholder="e.g., 1")

    search_clicked = st.button("🔎 Search Jobs")

    if search_clicked:
        if job_role.strip() and location.strip():
            results = search_jobs(job_role, location, experience_level, job_type, foundit_experience)

            st.markdown("## 🎯 Job Search Results")

            for job in results:
                platform = job["title"].split(":")[0].strip().lower()

                if platform == "linkedin":
                    icon = "🔵 <b style='color:#0e76a8;'>in LinkedIn</b>"
                    btn_color = "#0e76a8"
                elif platform == "naukri":
                    icon = "🏢 <b style='color:#ff5722;'>Naukri</b>"
                    btn_color = "#ff5722"
                elif "foundit" in platform:
                    icon = "🌐 <b style='color:#7c4dff;'>Foundit (Monster)</b>"
                    btn_color = "#7c4dff"
                else:
                    icon = f"📄 <b>{platform.title()}</b>"
                    btn_color = "#00c4cc"

                st.markdown(f"""
<div style="
    background-color:#1e1e1e;
    padding:20px;
    border-radius:15px;
    margin-bottom:20px;
    border-left: 5px solid {btn_color};
    box-shadow: 0 0 15px {btn_color};
">
    <div style="font-size:20px; margin-bottom:8px;">{icon}</div>
    <div style="color:#ffffff; font-size:17px; margin-bottom:15px;">
        {job['title'].split(':')[1].strip()}
    </div>
    <a href="{job['link']}" target="_blank" style="text-decoration:none;">
        <button style="
            background-color:{btn_color};
            color:white;
            padding:10px 15px;
            border:none;
            border-radius:8px;
            font-size:15px;
            cursor:pointer;
            box-shadow: 0 0 10px {btn_color};
        ">
            🚀 View Jobs on {platform.title()} &rarr;
        </button>
    </a>
</div>
""", unsafe_allow_html=True)
        else:
            st.warning("⚠️ Please enter both the Job Role and Location to perform the search.")


    # Inject Glowing CSS for Cards
    st.markdown("""
    <style>
    @keyframes glow {
        0% {
            box-shadow: 0 0 5px rgba(255,255,255,0.2);
        }
        50% {
            box-shadow: 0 0 20px rgba(0,255,255,0.6);
        }
        100% {
            box-shadow: 0 0 5px rgba(255,255,255,0.2);
        }
    }

    .company-card {
        background-color: #1e1e1e;
        color: #ffffff;
        border-radius: 15px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.3);
        transition: transform 0.3s ease-in-out, box-shadow 0.3s ease-in-out;
        cursor: pointer;
        text-decoration: none;
        display: block;
        animation: glow 3s infinite alternate;
    }

    .company-card:hover {
        transform: scale(1.03);
        box-shadow: 0 0 25px rgba(0, 255, 255, 0.7), 0 0 10px rgba(0, 255, 255, 0.5);
    }

    .pill {
        display: inline-block;
        background-color: #333;
        padding: 6px 12px;
        border-radius: 999px;
        margin: 4px 6px 0 0;
        font-size: 13px;
    }

    .title-header {
        color: white;
        font-size: 26px;
        margin-top: 40px;
        font-weight: bold;
    }

    .company-logo {
        font-size: 26px;
        margin-right: 8px;
    }

    .company-header {
        font-size: 22px;
        font-weight: bold;
        display: flex;
        align-items: center;
    }
    </style>
    """, unsafe_allow_html=True)


    # ---------- Featured Companies ----------
    st.markdown("### <div class='title-header'>🏢 Featured Companies</div>", unsafe_allow_html=True)

    selected_category = st.selectbox("📂 Browse Featured Companies By Category", ["All", "tech", "indian_tech", "global_corps"])
    companies_to_show = get_featured_companies() if selected_category == "All" else get_featured_companies(selected_category)

    for company in companies_to_show:
        category_tags = ''.join([f"<span class='pill'>{cat}</span>" for cat in company['categories']])
        st.markdown(f"""
        <a href="{company['careers_url']}" class="company-card" target="_blank">
            <div class="company-header">
                <span class="company-logo">{company.get('emoji', '🏢')}</span>
                {company['name']}
            </div>
            <p>{company['description']}</p>
            {category_tags}
        </a>
        """, unsafe_allow_html=True)

    # ---------- Market Insights ----------
    st.markdown("### <div class='title-header'>📈 Job Market Trends</div>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 🚀 Trending Skills")
        for skill in JOB_MARKET_INSIGHTS["trending_skills"]:
            st.markdown(f"""
            <div class="company-card">
                <h4>🔧 {skill['name']}</h4>
                <p>📈 Growth Rate: {skill['growth']}</p>
            </div>
            """, unsafe_allow_html=True)

    with col2:
        st.markdown("#### 🌍 Top Job Locations")
        for loc in JOB_MARKET_INSIGHTS["top_locations"]:
            st.markdown(f"""
            <div class="company-card">
                <h4>📍 {loc['name']}</h4>
                <p>💼 Openings: {loc['jobs']}</p>
            </div>
            """, unsafe_allow_html=True)

    # ---------- Salary Insights ----------
    st.markdown("### <div class='title-header'>💰 Salary Insights</div>", unsafe_allow_html=True)
    for role in JOB_MARKET_INSIGHTS["salary_insights"]:
        st.markdown(f"""
        <div class="company-card">
            <h4>💼 {role['role']}</h4>
            <p>📅 Experience: {role['experience']}</p>
            <p>💵 Salary Range: {role['range']}</p>
        </div>
        """, unsafe_allow_html=True)

with tab4:
    # CSS styles for header, buttons, and cards
    st.markdown("""
        <style>
        .header-box {
            background: linear-gradient(to right, #0f2027, #203a43, #2c5364);
            border: 2px solid #00c3ff;
            padding: 20px;
            border-radius: 15px;
            text-align: center;
            margin-bottom: 30px;
            box-shadow: 0 0 15px #00c3ff88;
        }

        .header-box h2 {
            font-size: 30px;
            color: #fff;
            margin: 0;
            font-weight: bold;
        }

        .glow-header {
            font-size: 22px;
            text-align: center;
            color: #00c3ff;
            text-shadow: 0 0 10px #00c3ff;
            margin-top: 10px;
            margin-bottom: 5px;
            font-weight: 600;
        }

        .stRadio > div {
            flex-direction: row !important;
            justify-content: center;
        }

        .stRadio label {
            background: #1a1a1a;
            border: 1px solid #00c3ff;
            color: #00c3ff;
            padding: 10px 16px;
            margin: 4px;
            border-radius: 8px;
            cursor: pointer;
            transition: all 0.3s ease;
        }

        .stRadio label:hover {
            background-color: #00c3ff33;
        }

        .stRadio input:checked + div > label {
            background-color: #00c3ff;
            color: black;
            font-weight: bold;
        }

        .card {
            background: linear-gradient(135deg, #0f2027, #203a43, #2c5364);
            border: 2px solid #00c3ff;
            border-radius: 15px;
            padding: 15px 20px;
            margin: 10px 0;
            box-shadow: 0 0 15px #00c3ff88;
            transition: transform 0.3s ease, box-shadow 0.3s ease;
        }

        .card:hover {
            transform: scale(1.02);
            box-shadow: 0 0 25px #00c3ffcc;
        }

        .card a {
            color: #00c3ff;
            font-weight: bold;
            font-size: 16px;
            text-decoration: none;
        }

        .card a:hover {
            color: #ffffff;
            text-decoration: underline;
        }
        </style>
    """, unsafe_allow_html=True)

    # Header
    st.markdown("""
        <div class="header-box">
            <h2>📚 Recommended Learning Hub</h2>
        </div>
    """, unsafe_allow_html=True)

    # Subheader
    st.markdown('<div class="glow-header">🎓 Explore Career Resources</div>', unsafe_allow_html=True)
    st.markdown("<p style='text-align:center; color:#ccc;'>Curated courses and videos for your career growth, resume tips, and interview success.</p>", unsafe_allow_html=True)

    # Enhanced label above radio buttons
    st.markdown("""
        <div style="text-align:center; margin-top: 20px; margin-bottom: 10px;">
            <span style="color: #00c3ff; font-weight: bold; font-size: 20px; text-shadow: 0 0 10px #00c3ff;">
                🧭 Choose Your Learning Path
            </span>
        </div>
    """, unsafe_allow_html=True)

    # Stylish radio buttons
    page = st.radio(
        " ",
        ["Courses by Role", "Resume Videos", "Interview Videos"],
        horizontal=True,
        key="page_selection"
    )

    if page == "Courses by Role":
        st.subheader("Courses by Career Role")
        category = st.selectbox(
            "Select Career Category",
            options=list(COURSES_BY_CATEGORY.keys()),
            key="category_selection"
        )
        if category:
            roles = list(COURSES_BY_CATEGORY[category].keys())
            role = st.selectbox(
                "Select Role / Job Title",
                options=roles,
                key="role_selection"
            )
            if role:
                st.subheader(f"Courses for {role} in {category}:")
                courses = get_courses_for_role(category, role)
                if courses:
                    for title, url in courses:
                        st.markdown(f"""
                            <div class="card">
                                <a href="{url}" target="_blank">🔗 {title}</a>
                            </div>
                        """, unsafe_allow_html=True)
                else:
                    st.info("No courses found for this role.")

    elif page == "Resume Videos":
        st.subheader("📄 Resume Writing Videos")
        categories = list(RESUME_VIDEOS.keys())
        selected_cat = st.selectbox(
            "Select Resume Video Category",
            options=categories,
            key="resume_vid_cat"
        )
        if selected_cat:
            st.subheader(f"{selected_cat}")
            videos = RESUME_VIDEOS[selected_cat]
            cols = st.columns(2)
            for idx, (title, url) in enumerate(videos):
                with cols[idx % 2]:
                    st.markdown(f"**{title}**")
                    st.video(url)

    elif page == "Interview Videos":
        st.subheader("🗣️ Interview Preparation Videos")
        categories = list(INTERVIEW_VIDEOS.keys())
        selected_cat = st.selectbox(
            "Select Interview Video Category",
            options=categories,
            key="interview_vid_cat"
        )
        if selected_cat:
            st.subheader(f"{selected_cat}")
            videos = INTERVIEW_VIDEOS[selected_cat]
            cols = st.columns(2)
            for idx, (title, url) in enumerate(videos):
                with cols[idx % 2]:
                    st.markdown(f"**{title}**")
                    st.video(url)
with tab5:
    import sqlite3
    import pandas as pd
    from db_manager import get_top_domains_by_score, delete_candidate_by_id

    st.markdown("## 🛡️ <span style='color:#336699;'>Admin Database Panel</span>", unsafe_allow_html=True)
    st.markdown("<hr style='border-top: 2px solid #bbb;'>", unsafe_allow_html=True)

    # 🔐 Admin Login
    if not st.session_state.get("admin_logged_in", False):
        st.warning("🔒 Admin access required.")
        password = st.text_input("Enter Admin Password", type="password")
        if st.button("Login"):
            if password == "lexiadmin123":  # ⚠️ Replace with secure method in prod
                st.session_state.admin_logged_in = True
                st.success("✅ Logged in successfully.")
                st.rerun()
            else:
                st.error("❌ Incorrect password.")
        st.stop()

    # ✅ Admin Logged In
    st.success("✅ You are logged in as admin.")

    # 🔄 Manual Refresh Button
    if st.button("🔄 Refresh Data"):
        st.rerun()

    # 📄 Load Data from DB
    conn = sqlite3.connect("resume_data.db")
    df = pd.read_sql_query("SELECT * FROM candidates ORDER BY timestamp DESC", conn)

    # 🔍 Search by Candidate Name
    search = st.text_input("🔍 Search Candidate by Name")
    if search:
        df = df[df["candidate_name"].str.contains(search, case=False, na=False)]

    # 📋 Display DataFrame
    if df.empty:
        st.info("ℹ️ No candidate data found.")
    else:
        st.markdown("### 📋 Candidate Submissions")
        st.dataframe(df, use_container_width=True)

        # 📥 Download CSV Export
        st.download_button(
            label="📥 Download CSV",
            data=df.to_csv(index=False),
            file_name="candidates_export.csv",
            mime="text/csv",
            use_container_width=True
        )

        # 🗑️ Delete Candidate
        st.markdown("### 🗑️ Delete Candidate Entry")
        delete_id = st.number_input("Enter Candidate ID to Delete", min_value=1, step=1)
        if st.button("🗑 Delete Candidate"):
            if delete_id in df["id"].values:
                delete_candidate_by_id(delete_id)
                st.success(f"✅ Candidate with ID {delete_id} deleted.")
                st.rerun()
            else:
                st.warning("⚠️ ID not found in current table.")

    # 📊 Top 5 Domains by ATS Score
    st.markdown("### 📊 Top 5 Domains by ATS Score")
    top_domains = get_top_domains_by_score()
    if top_domains:
        for domain, avg_score, count in top_domains:
            st.info(f"📁 **{domain}** — Average ATS: {avg_score:.2f} ({count} resumes)")
    else:
        st.info("ℹ️ No domain analytics available.")

# 1. Display existing chat history
if "memory" in st.session_state:
    history = st.session_state.memory.load_memory_variables({}).get("chat_history", [])
    for msg in history:
        with st.chat_message("user" if msg.type == "human" else "assistant"):
            st.markdown(msg.content)

# 2. Wait for user input
user_input = st.chat_input("Ask LEXIBOT anything...")

# 3. Only call chain when user submits new input
if user_input:
    # Show user message
    with st.chat_message("user"):
        st.markdown(user_input)

    try:
        # 🧠 Only call chain ONCE
        response = st.session_state.chain.invoke({
            "question": user_input,
            "chat_history": st.session_state.memory.chat_memory.messages
        })
        answer = response.get("answer", "❌ No answer found.")
    except Exception as e:
        answer = f"⚠️ Error: {str(e)}"

    # Show assistant reply
    with st.chat_message("assistant"):
        st.markdown(answer)

    # Save interaction to memory
    st.session_state.memory.save_context({"input": user_input}, {"output": answer})
